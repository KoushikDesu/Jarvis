import os
import json

class DocumentParser:
    def __init__(self):
        # We lazy-load heavy packages to keep server start times fast
        self.ocr_reader = None

    def parse_file(self, file_path):
        """Parse file based on extension and return extracted text."""
        if not os.path.exists(file_path):
            return f"Error: File '{os.path.basename(file_path)}' not found."

        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                return self.parse_pdf(file_path)
            elif ext in ['.doc', '.docx']:
                return self.parse_docx(file_path)
            elif ext == '.json':
                return self.parse_json(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
                return self.parse_image(ocr_fallback=True, file_path=file_path)
            elif ext in ['.txt', '.md', '.py', '.js', '.css', '.html', '.csv']:
                return self.parse_text_file(file_path)
            else:
                return f"Unsupported file type '{ext}'. Displaying metadata: Size = {os.path.getsize(file_path)} bytes."
        except Exception as e:
            return f"Error parsing {os.path.basename(file_path)}: {str(e)}"

    def parse_text_file(self, file_path):
        """Simple plain text files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return content

    def parse_pdf(self, file_path):
        """Extract text from PDF file using pypdf."""
        try:
            import pypdf
        except ImportError:
            return "Error: 'pypdf' package is not installed. Run 'pip install pypdf' to enable PDF parsing."

        reader = pypdf.PdfReader(file_path)
        text_content = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_content.append(f"--- Page {idx+1} ---\n{text}")
        
        return "\n".join(text_content) if text_content else "PDF was parsed successfully, but no text could be extracted (it might be scanned). Try using OCR on the PDF pages."

    def parse_docx(self, file_path):
        """Extract text from Word Document using python-docx."""
        try:
            import docx
        except ImportError:
            return "Error: 'python-docx' package is not installed. Run 'pip install python-docx' to enable Word document parsing."

        doc = docx.Document(file_path)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        # Also parse tables inside word doc
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_content.append(" | ".join(row_data))
                
        return "\n".join(text_content)

    def parse_json(self, file_path):
        """Parse and prettify JSON files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            data = json.load(f)
        return json.dumps(data, indent=2, ensure_ascii=False)

    def parse_image(self, file_path, ocr_fallback=True):
        """Perform OCR on images locally using easyocr."""
        if not ocr_fallback:
            return "[IMAGE DATA SENT TO GEMINI API]"
            
        try:
            import easyocr
            import numpy as np
            from PIL import Image
        except ImportError:
            return "Error: OCR libraries not installed. Run 'pip install easyocr pillow numpy' to enable local image text extraction."

        try:
            # Initialize easyocr reader on GPU if CUDA is available, else CPU
            if self.ocr_reader is None:
                print("Initializing EasyOCR reader (this may take a few seconds on first run)...")
                # English, Hindi, and Telugu OCR capability
                self.ocr_reader = easyocr.Reader(['en', 'hi', 'te'], gpu=True)

            print(f"Performing OCR on image: {os.path.basename(file_path)}")
            # Read text from image file path
            results = self.ocr_reader.readtext(file_path)
            
            # Sort results by their vertical position, then horizontal position (reading order)
            results = sorted(results, key=lambda x: (x[0][0][1], x[0][0][0]))
            
            extracted_text = []
            for bbox, text, confidence in results:
                if confidence > 0.15: # Filter out noise
                    extracted_text.append(text)
                    
            if not extracted_text:
                return "OCR completed: No text detected in this image."
                
            return "\n".join(extracted_text)
        except Exception as e:
            return f"OCR Extraction failed: {str(e)}"

if __name__ == "__main__":
    # Test stub
    parser = DocumentParser()
    print("DocumentParser initialized successfully.")
